from docx import Document
from app.utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from paragraphs and tables in a .docx file."""
    logger.info(f"Extracting text from DOCX: {file_path}")
    try:
        doc = Document(file_path)
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Tables are not in .paragraphs — handle separately
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    lines.append(" | ".join(row_text))

        full_text = "\n".join(lines)
        logger.info(f"DOCX extraction complete. Characters: {len(full_text)}")
        return full_text
    except FileNotFoundError:
        logger.error(f"DOCX not found: {file_path}")
        return ""
    except Exception as e:
        logger.error(f"Failed to extract DOCX: {file_path} | {e}")
        return ""
